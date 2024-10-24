import pandas as pd
import os
import json
import logging
from datetime import datetime
from docx import Document
from decouple import config
from typing import List

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_ollama import OllamaLLM
from langchain_ollama import OllamaEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.prompts import PromptTemplate
from langchain.callbacks.manager import CallbackManager
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
from langchain_core.output_parsers import StrOutputParser
from langchain_core.output_parsers import BaseOutputParser

import faiss
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_community.document_transformers import LongContextReorder
from langchain.retrievers.multi_query import MultiQueryRetriever

import pdf_a_texto
from logging_config import setup_logging


# Definir el directorio base de la solución
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# Configuracion del logging
log_file_path = setup_logging()
logger = logging.getLogger(__name__)


try:

    # Leer rutas a plantillas
    carpeta_entrada = config("INPUT_FOLDER")
    carpeta_salida = config("OUTPUT_FOLDER")
    plantilla_excel = config("PROMPTS_TEMPLATE")
    plantilla_word = config("FORMAT_TEMPLATE")
    prompt_sistema = config("SYSTEM_PROMPT")
    plantilla_contexto = config("CONTEXT_TEMPLATE")
    plantilla_multi = config("MULTIQUERY_TEMPLATE")

    logger.info("Configuración y rutas cargadas correctamente.")

    # Inicializar la base de datos vectorial
    embeddings_model = config("EMBEDDING_MODEL")
    local_embeddings = OllamaEmbeddings(model=embeddings_model, base_url="http://host.docker.internal:11434")
    logger.info(f"Modelo embeddings: {embeddings_model}")
    
        
    index = faiss.IndexFlatL2(len(local_embeddings.embed_query("hola mundo")))
    vectorstore = FAISS(
        embedding_function=local_embeddings,
        index=index,
        docstore=InMemoryDocstore(),
        index_to_docstore_id={},
        normalize_L2=True
    )

    # Indicamos el modelo LLM
    with open(prompt_sistema, 'r', encoding='utf-8') as system_p:
        SYSTEM_TEMPLATE = system_p.read()
    
    model = OllamaLLM(
        base_url="http://host.docker.internal:11434",
        system=SYSTEM_TEMPLATE,
        model=config("DEFAULT_MODEL"), 
        callback_manager=CallbackManager([StreamingStdOutCallbackHandler()]),
        temperature=config("TEMPERATURE"),
        top_p=config("TOP_P"),
    )
    logger.info(f"Modelo de LLM: {config('DEFAULT_MODEL')} con temperature: {config('TEMPERATURE')} y top_p: {config('TOP_P')}")


    # Procesar cada documento en la carpeta de entrada
    for filename in os.listdir(carpeta_entrada):
        filepath = os.path.join(carpeta_entrada, filename)
        if filename.endswith(".pdf"):
            logger.info(f"Procesando archivo: {filepath}")
            data = pdf_a_texto.procesar_pdf(filepath, config("RESOLUCION_OCR")) # OCR si es necesario
            logger.info("Archivo PDF cargado correctamente.")

            # Dividir el archivo en chunks
            text_splitter = RecursiveCharacterTextSplitter(config("CHUNK_SIZE"), config("CHUNK_OVERLAP"))
            all_splits = text_splitter.split_text(data)
            logger.info(f"Documento {filename} dividido en chunks correctamente.")

            # Añadir los chunks al índice FAISS
            vectorstore.add_texts(texts=all_splits, embedding=local_embeddings)
            logger.info(f"Fragmentos de {filename} añadidos a la base de datos vectorial.")

    logger.info("Base de datos vectorial completada correctamente.")

    # Leer plantillas RAG
    with open(plantilla_multi, 'r', encoding='utf-8') as multi:
        MULTI_TEMPLATE = multi.read()

    multi_prompt = PromptTemplate.from_template(MULTI_TEMPLATE)
    logger.info("Leida la plantilla para hacer multiquery de la pregunta del usuario.")

    with open(plantilla_contexto, 'r', encoding='utf-8') as contexto:
        RAG_TEMPLATE = contexto.read()
   
    rag_prompt = ChatPromptTemplate.from_template(RAG_TEMPLATE)
    logger.info("Leida la plantilla para RAG.")    
    

    # Funcion para splitear la salida por lineas
    class LineListOutputParser(BaseOutputParser[List[str]]):
        """Output parser for a list of lines."""

        def parse(self, text: str) -> List[str]:
            lines = text.strip().split("\n")
            return list(filter(None, lines))  # Remove empty lines

    output_parser = LineListOutputParser()

    def format_docs(docs):
        return "\n\n".join(doc.page_content for doc in docs)

    # Chain para hacer multiquery
    multi_chain = ( 
        {
            "question": lambda x: x["question"]
        }
        | multi_prompt
        | model
        | output_parser
    )

    # Chain para hacer RAG
    qa_chain = (
        {
            "question": lambda x: x["question"],  # input query
            "context": lambda x: format_docs(x["context"]),  # context
        }
        | rag_prompt
        | model
        | StrOutputParser()
    )
    logger.info("Chain de RAG creada correctamente.")

    # Definimos el retriever
    retriever = MultiQueryRetriever.from_llm(
        retriever=vectorstore.as_retriever(
            search_type="mmr",
            search_kwargs={"k": int(config("K"))}
        ),
        llm=model,
        parser_key="lines"
    )

    # LEER PROMPTS Y CARGAR PLANTILLAS
    queries_df = pd.read_excel(plantilla_excel)
    logger.info(f"La plantilla con los prompts es: {plantilla_excel}")

    # Cargar el documento .docx
    doc = Document(plantilla_word)
    logger.info("Plantilla Word cargada correctamente.")
    output_file_path_docx = os.path.join(carpeta_salida, f"CU1_MEMORIA_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")

    
    # EJECUCION
    for _, row in queries_df.iterrows():
        question = row['PROMPT']
        marcador = row['MARCADOR']
        logger.info(f"Procesando PROMPT: {question}")

        if pd.isna(question) or question.strip() == "":
            logger.warning("Pregunta vacía o nula, se omite.")
            continue

        try:
            # Buscamos trozos relevantes en el contexto utilizando multiquery
            multiquery = multi_chain.invoke({"question": question})
            logger.info(f"MULTIQUERY: {multiquery}")
            docs = retriever.invoke(multiquery)
            # Funcion para reordenar los trozos relevantes del contexto
            reordering = LongContextReorder()
            reordered_docs = reordering.transform_documents(docs)
            logger.info(f"RETRIEVED DOCS: {reordered_docs}")
            # Respuesta a la pregunta original
            respuesta = qa_chain.invoke({"context": reordered_docs, "question": question})
            logger.info(f"RESPUESTA: {respuesta}")
        except Exception as e:
            logger.error(f"Error al procesar la consulta: {str(e)}")
            respuesta = f'Error al procesar la consulta: {str(e)}'

        # Buscar el marcador en el documento y reemplazarlo con la respuesta
        found = False
        for para in doc.paragraphs:
            if marcador in para.text:
                para.text = para.text.replace(marcador, respuesta)
                found = True
                break
        if not found:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if marcador in cell.text:
                            cell.text = cell.text.replace(marcador, respuesta)
                            found = True
                            break
                    if found:
                        break
                if found:
                    break
        if not found:
            logger.warning(f"Marcador '{marcador}' no encontrado en el documento.")
        doc.save(output_file_path_docx)
    
    
    logger.info(f"Se ha generado la memoria: {output_file_path_docx}")

except Exception as e:
    logger.critical(f"Error crítico durante la ejecución del script: {str(e)}", exc_info=True)
    print(f"Ha ocurrido un error crítico. Revisa el log para más detalles: {log_file_path}")


