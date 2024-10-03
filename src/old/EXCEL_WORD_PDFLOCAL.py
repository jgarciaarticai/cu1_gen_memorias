import pandas as pd
from langchain import PromptTemplate
from langchain.llms import Ollama
from langchain.callbacks.manager import CallbackManager
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
from PyPDF2 import PdfReader
import sys
import os
import json
from dotenv import load_dotenv
from datetime import datetime
from docx import Document

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# Cargar configuración general desde config.json
config_path = os.path.join(BASE_DIR, 'config', 'config.json')
with open(config_path, 'r', encoding='utf-8') as config_file:
    config = json.load(config_file)

archivo_entrada = config["input_file"]
carpeta_salida = config["output_folder"]
modelo = config["default_model"]
plantilla_excel = config["prompts_template"]
plantilla_word = config["format_template"]
plantilla_contexto = config["context_template"]

class SuppressStdout:
    def __enter__(self):
        self._original_stdout = sys.stdout
        self._original_stderr = sys.stderr
        sys.stdout = open(os.devnull, 'w')
        sys.stderr = open(os.devnull, 'w')

    def __exit__(self, exc_type, exc_val, exc_tb):
        sys.stdout.close()
        sys.stdout = self._original_stdout
        sys.stderr = self._original_stderr

# Leer el contenido del PDF
def leer_pdf(archivo_pdf):
    reader = PdfReader(archivo_pdf)
    texto = ""
    for page in reader.pages:
        texto += page.extract_text()
    return texto

contenido_pdf = leer_pdf(archivo_entrada)
print(f"El archivo que se va a leer es: {archivo_entrada}")

# Leer los prompts de la plantilla    
queries_df = pd.read_excel(plantilla_excel)
print(f"La plantilla con los prompts es: {plantilla_excel}")

if 'PROMPT' not in queries_df.columns:
    raise ValueError("El archivo Excel debe contener una columna llamada 'PROMPT'.")

# Plantilla de prompt context
with open(plantilla_contexto, 'r') as contexto:
    template = contexto.read()

QA_CHAIN_PROMPT = PromptTemplate(
    input_variables=["context", "question"],
    template=template,
)

# Configuración del LLM
llm = Ollama(model=modelo, callback_manager=CallbackManager([StreamingStdOutCallbackHandler()]))

# Cargar el documento .docx
doc = Document(plantilla_word)

# Preparar el fichero de salida
txt_file_path = carpeta_salida + "/memoriaSINRAG_" + datetime.now().strftime("%Y%m%d_%H%M%S") + ".txt"


# Procesar las preguntas y escribir las respuestas en el documento
with open(txt_file_path, 'w', encoding='utf-8') as txt_file:
    for _, row in queries_df.iterrows():
        prompt = row['PROMPT']
        marcador = row['MARCADOR']
        print(f"\n\nPROMPT: {prompt}")
        if pd.isna(prompt) or prompt.strip() == "":
            continue

        # Usar el contenido del PDF como contexto
        context = contenido_pdf

        # Generar el texto del prompt
        prompt_text = QA_CHAIN_PROMPT.format(context=context, question=prompt)

        # Preguntar al modelo
        try:
            result = llm(prompt_text)
            answer = result
        except Exception as e:
            answer = f'Error al procesar la consulta: {str(e)}'
    
        # Escribir pregunta y respuesta en un fichero .txt
        txt_file.write(f"Pregunta: {prompt}\nRespuesta: {answer}\n\n")
    

        # Buscar el marcador en el documento y reemplazarlo con la respuesta
        found = False
        for para in doc.paragraphs:
            if marcador in para.text:
                para.text = para.text.replace(marcador, answer)
                found = True
                break
        if not found:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if marcador in cell.text:
                            cell.text = cell.text.replace(marcador, answer)
                            found = True
                            break
                    if found:
                        break
                if found:
                    break
        if not found:
            print(f"Marcador '{marcador}' no encontrado en el documento.")

# Guardar el documento con las respuestas insertadas
output_file_path = os.path.join(carpeta_salida, "CU1_MEMORIA_" + datetime.now().strftime("%Y%m%d_%H%M%S") + ".docx")
doc.save(output_file_path)

print(f"DOCUMENTO GENERADO: {output_file_path}")
