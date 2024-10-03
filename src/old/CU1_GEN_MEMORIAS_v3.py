import pandas as pd
import os
import json
import logging
from datetime import datetime
from docx import Document

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_ollama import OllamaEmbeddings
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.prompts import PromptTemplate
from langchain.chains import RetrievalQA
from langchain_community.llms import Ollama
from langchain.callbacks.manager import CallbackManager
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler

# Definir el directorio base de la solución
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# Configuración de logging para almacenar logs en la carpeta de logs en el directorio raíz de la solución
log_dir = os.path.join(BASE_DIR, 'logs')
os.makedirs(log_dir, exist_ok=True)  # Crear el directorio si no existe
log_file_path = os.path.join(log_dir, f"gen_memorias_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log")

logging.basicConfig(
    filename=log_file_path,
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

try:
    # Cargar configuración general desde config.json
    config_path = os.path.join(BASE_DIR, 'config', 'config.json')
    with open(config_path, 'r', encoding='utf-8') as config_file:
        config = json.load(config_file)
    logger.info("Configuración cargada correctamente.")

    # Leer rutas a plantillas
    archivo_entrada = config["input_file"]
    carpeta_salida = config["output_folder"]
    plantilla_excel = config["prompts_template"]
    plantilla_word = config["format_template"]
    plantilla_contexto = config["context_template"]
    plantilla_refinamiento = config["refine_template"]

    # CARGA Y SPLIT DEL DOCUMENTO
    logger.info(f"El archivo que se va a leer es: {archivo_entrada}")
    loader = PyPDFLoader(archivo_entrada)
    data = loader.load()
    logger.info("Archivo PDF cargado correctamente.")

    # Dividir el archivo en chunks
    text_splitter = RecursiveCharacterTextSplitter(config["chunk_size"], config["chunk_overlap"])
    all_splits = text_splitter.split_documents(data)
    logger.info("Documento dividido en chunks correctamente.")

    # Inicializar la base de datos vectorial y pasar los chunks
    embeddings_model = config["embedding_model"]
    
    local_embeddings = OllamaEmbeddings(model=embeddings_model)
    logger.info(f"Modelo embeddings: {embeddings_model}")
    vectorstore = Chroma.from_documents(documents=all_splits, embedding=local_embeddings)
    logger.info(f"Base de datos vectorial inicializada correctamente.")

    # Indicamos el modelo
    model=config["default_model"]
    temperature = config["temperature"]
    top_p=config["top_p"]

    model = Ollama(
        model=model, 
        callback_manager=CallbackManager([StreamingStdOutCallbackHandler()]),
        temperature=temperature,
        top_p=top_p,
    )
    logger.info(f"temperature: {temperature} y top_p: {top_p}")

    # RAG
    with open(plantilla_contexto, 'r', encoding='utf-8') as contexto:
        RAG_TEMPLATE = contexto.read()

    rag_prompt = PromptTemplate(
        input_variables=["context", "question"],
        template=RAG_TEMPLATE
    )

    k = int(config["k"])
    chain_type=config["chain_type"]
    
    qa_chain = RetrievalQA.from_chain_type(
        model,
        retriever=vectorstore.as_retriever(search_kwargs={"k": k}),
        chain_type=chain_type,
        chain_type_kwargs={"prompt": rag_prompt}
    )
    logger.info(f"Chain de RAG creada correctamente. Chain type: {chain_type}")

    # LEER PROMPTS Y CARGAR PLANTILLAS
    queries_df = pd.read_excel(plantilla_excel)
    logger.info(f"La plantilla con los prompts es: {plantilla_excel}")

    # Cargar el documento .docx
    doc = Document(plantilla_word)
    logger.info("Plantilla Word cargada correctamente.")

    # Preparar el fichero de salida
    output_file_path = os.path.join(carpeta_salida, f"memoria_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt")
    logger.info(f"El fichero de salida es: {output_file_path}")

    # EJECUCION
    with open(output_file_path, 'w', encoding='utf-8') as output_file:
        for _, row in queries_df.iterrows():
            question = row['PROMPT']
            marcador = row['MARCADOR']
            logger.info(f"Procesando PROMPT: {question}")

            if pd.isna(question) or question.strip() == "":
                logger.warning("Pregunta vacía o nula, se omite.")
                continue

            try:
                respuesta = qa_chain.invoke({"query": question})
                answer = respuesta.get('result', 'No se encontró una respuesta.')
                logger.info(f"RESPUESTA: {respuesta}")
            except Exception as e:
                logger.error(f"Error al procesar la consulta: {str(e)}")
                answer = f'Error al procesar la consulta: {str(e)}'
            
            output_file.write(f"Pregunta: {question}\nRespuesta: {answer}\n\n")
            
            # Buscar el marcador en el documento y reemplazarlo con la respuesta
            found = False
            for para in doc.paragraphs:
                if marcador in para.text:
                    para.text = para.text.replace(marcador, answer)
                    found = True
                    break
            if not found:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if marcador in cell.text:
                                cell.text = cell.text.replace(marcador, answer)
                                found = True
                                break
                        if found:
                            break
                    if found:
                        break
            if not found:
                logger.warning(f"Marcador '{marcador}' no encontrado en el documento.")

    output_file_path_docx = os.path.join(carpeta_salida, f"CU1_MEMORIA_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(output_file_path_docx)
    logger.info(f"Se ha generado la memoria: {output_file_path_docx}")

except Exception as e:
    logger.critical(f"Error crítico durante la ejecución del script: {str(e)}", exc_info=True)
    print(f"Ha ocurrido un error crítico. Revisa el log para más detalles: {log_file_path}")
