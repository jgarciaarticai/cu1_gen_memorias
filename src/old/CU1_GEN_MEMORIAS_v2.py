import pandas as pd
import os
import json
from datetime import datetime
from docx import Document

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_ollama import OllamaEmbeddings
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain.chains import RetrievalQA
from langchain_community.llms import Ollama
from langchain.callbacks.manager import CallbackManager
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
from langchain.chains import LLMChain, RefineDocumentsChain
from langchain_core.runnables import RunnablePassthrough


BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


# Cargar configuración general desde config.json
config_path = os.path.join(BASE_DIR, 'config', 'config.json')
with open(config_path, 'r', encoding='utf-8') as config_file:
    config = json.load(config_file)

# Leer rutas a plantillas
archivo_entrada = config["input_file"]
carpeta_salida = config["output_folder"]
plantilla_excel = config["prompts_template"]
plantilla_word = config["format_template"]
plantilla_contexto = config["context_template"]
plantilla_refinamiento = config["refine_template"]


### CARGA Y SPLIT DEL DOCUMENTO
# Cargar el PDF
print(f"El archivo que se va a leer es: {archivo_entrada}")
loader = PyPDFLoader(archivo_entrada)
data = loader.load()

# Dividir el archivo en chunks
text_splitter = RecursiveCharacterTextSplitter(config["chunk_size"], config["chunk_overlap"])
all_splits = text_splitter.split_documents(data)

# Inicializar la base de datos vectorial y pasar los chunks
local_embeddings = OllamaEmbeddings(model=config["embedding_model"])
vectorstore = Chroma.from_documents(documents=all_splits, embedding=local_embeddings)
retriever = vectorstore.as_retriever(search_kwargs={"k": int(config["k"])})

# Indicamos el modelo
model = Ollama(
    model=config["default_model"], 
    callback_manager=CallbackManager([StreamingStdOutCallbackHandler()]),
    temperature=config["temperature"],
    top_p=config["top_p"],
)


### RAG
# Se crea la plantilla para RAG a partir de la plantilla de contexto
with open(plantilla_contexto, 'r', encoding='utf-8') as contexto:
    RAG_TEMPLATE = contexto.read()

rag_prompt = PromptTemplate(
    input_variables=["context", "question"],
    template=RAG_TEMPLATE
)

# Se crea la chain inicial
qa_chain = RetrievalQA.from_chain_type(
    model,
    retriever=retriever,
    chain_type=config["chain_type"],
    chain_type_kwargs={"prompt": rag_prompt}
)


### REFINADO DEL RAG
# Se crea la plantilla para realizar el refinado de las respuestas iniciales
with open(plantilla_refinamiento, 'r', encoding='utf-8') as refinamiento:
    REFINE_TEMPLATE = refinamiento.read()

# Prompt de refinamiento
refine_prompt = PromptTemplate(
    input_variables=["existing_answer", "context"],
    template=REFINE_TEMPLATE,
)
# Chain inicial que se usa para el refinamiento
initial_llm_chain = LLMChain(
    llm=model, 
    prompt=rag_prompt
)

# Chain de refinamiento
refine_llm_chain = LLMChain(
    llm=model, 
    prompt=refine_prompt
)

# Se va a depreciar LLMChain, habria que poner esto:
# initial_llm_chain = rag_prompt | model | StrOutputParser()
#refine_llm_chain = refine_prompt | model | StrOutputParser()

chain = RefineDocumentsChain(
    initial_llm_chain=initial_llm_chain,
    refine_llm_chain=refine_llm_chain,
    initial_response_name="existing_answer",
    document_variable_name="context",
    input_key="query",
    output_key="result",
)

rag_refine_chain = (
    {"context": retriever, "query": RunnablePassthrough()}
    | chain
    | StrOutputParser()
)


# LEER PROMPTS Y CARGAR PLANTILLAS
# Leer los prompts de la plantilla Excel
queries_df = pd.read_excel(plantilla_excel)
print(f"La plantilla con los prompts es: {plantilla_excel}")

# Cargar el documento .docx
doc = Document(plantilla_word)

# Preparar el fichero de salida
output_file_path = carpeta_salida + "/memoria_" + datetime.now().strftime("%Y%m%d_%H%M%S") + ".txt"
print(f"El fichero de salida es: {output_file_path}")


# EJECUCION
# Recorrer la plantilla Excel para hacer las preguntas al modelo
with open(output_file_path, 'w', encoding='utf-8') as output_file:
    for _, row in queries_df.iterrows():
        question = row['PROMPT']
        marcador = row['MARCADOR']
        print(f"\n\nPROMPT: {question}")
        if pd.isna(question) or question.strip() == "":
            continue

        # Preguntar al modelo
        try:
            respuesta_inicial = qa_chain.invoke({"query": question})
            respuesta_refinada = rag_refine_chain.invoke({"query": question, "existing_answer" : respuesta_inicial['result']})
            answer = respuesta_refinada.get('result', 'No se encontró una respuesta.')
   
        except Exception as e:
            answer = f'Error al procesar la consulta: {str(e)}'
            
        # Escribir pregunta y respuesta en el fichero de salida
        output_file.write(f"Pregunta: {question}\nRespuesta: {answer}\n\n")
        
        # Buscar el marcador en el documento y reemplazarlo con la respuesta
        found = False
        for para in doc.paragraphs:
            if marcador in para.text:
                para.text = para.text.replace(marcador, answer)
                found = True
                break
        if not found:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if marcador in cell.text:
                            cell.text = cell.text.replace(marcador, answer)
                            found = True
                            break
                    if found:
                        break
                if found:
                    break
        if not found:
            print(f"Marcador '{marcador}' no encontrado en el documento.")

          
# Guardar el documento con las respuestas insertadas
output_file_path = carpeta_salida + "/CU1_MEMORIA_" + datetime.now().strftime("%Y%m%d_%H%M%S") + ".docx"
doc.save(output_file_path)

print(f"\nSe ha generado la memoria: {output_file_path}")